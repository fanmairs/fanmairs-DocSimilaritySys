from __future__ import annotations

import csv
import json
import math
import re
import shutil
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, List, Optional


ROOT = Path(__file__).resolve().parents[1]
DATA_DIR = ROOT / "data"
OUTPUT_DIR = ROOT / "experiments" / "threshold_sensitivity_samples_final"
TARGET_DIR = OUTPUT_DIR / "samples" / "targets"
REFERENCE_DIR = OUTPUT_DIR / "samples" / "references"


@dataclass(frozen=True)
class SourceDoc:
    filename: str
    domain: str


@dataclass
class ExtractedDoc:
    source: SourceDoc
    path: Path
    paragraphs: List[str]
    char_count: int


@dataclass(frozen=True)
class SamplePair:
    case_id: str
    sample_type: str
    label: int
    target_text: str
    reference_text: str
    target_source: str
    reference_source: str
    domain: str
    notes: str


SOURCE_DOCS = [
    SourceDoc("基于文本相似度的科研项目查重算法研究及应用_吕云峰.pdf", "text_similarity"),
    SourceDoc("基于文本语义增强的学术论文相似度检测方法研究及应用_张旭.pdf", "text_similarity"),
    SourceDoc("多特征融合的短文本语义相似度计算方法研究_张凝.pdf", "text_similarity"),
    SourceDoc("深度学习框架下系统查重的研究与应用_雷子旺.pdf", "text_similarity"),
    SourceDoc("基于小样本学习的文档查重系统的设计与实现_刘宏更.pdf", "text_similarity"),
    SourceDoc("基于混合检测策略的作业查重系统研究与实现_胡布焕.pdf", "text_similarity"),
    SourceDoc("基于语义神经网络的文档查重系统设计与实现_戴文丽.pdf", "text_similarity"),
    SourceDoc("基于TF-IDF和机器学习的文本向量化与分类研究_石慧.pdf", "text_similarity"),
    SourceDoc("基于Transformer的多尺度特征融合的小目标检测研究_王德玉.pdf", "computer_vision"),
    SourceDoc("基于改进YOLOv11的小目标行人检测与多目标跟踪方法研究_邢世洋.pdf", "computer_vision"),
    SourceDoc("基于改进YOLO的绝缘子缺陷检测算法研究_刘云.pdf", "computer_vision"),
    SourceDoc("基于YOLOv5的轻量级小目标检测算法与交通流量统计_黄益鑫.pdf", "computer_vision"),
    SourceDoc("基于计算机视觉的水果自动分级技术研究_苗天强.pdf", "computer_vision"),
    SourceDoc("基于卷积神经网络的人脸表情识别与应用_王以诺.pdf", "computer_vision"),
    SourceDoc("C县农产品物流园区选址及功能区布局规划研究_张明慧.pdf", "agri_logistics"),
    SourceDoc("京津冀农产品物流需求预测与物流网络构建研究_薛晓博.pdf", "agri_logistics"),
    SourceDoc("FCH公司生鲜农产品配送路径优化研究_徐小雨.pdf", "agri_logistics"),
    SourceDoc("双碳目标下生鲜农产品冷链配送路径优化研究_牛忍.pdf", "agri_logistics"),
    SourceDoc("农业全产业链发展水平评价研究_徐靖.pdf", "agri_logistics"),
    SourceDoc("政府规制下生鲜农产品供应链的区块链溯源决策研究_蒋凤娇.pdf", "agri_logistics"),
    SourceDoc("“新事实-新规范”路径下虚拟财产的解释_陈亦洲.pdf", "law"),
    SourceDoc("个人数字遗产继承的法律保护研究_李苗苗.pdf", "law"),
    SourceDoc("数字遗产继承问题研究_李欣.pdf", "law"),
    SourceDoc("平台处罚的行政法规制_吴树贞.pdf", "law"),
    SourceDoc("网络暴力治理视角下平台主体责任研究_刘鹏飞.pdf", "law"),
    SourceDoc("影评类短视频侵权纠纷中合理使用认定研究_连宇馨.pdf", "law"),
    SourceDoc("AI协同下高中数学实验教学模式构建与实践研究_白雪.pdf", "education"),
    SourceDoc("核心素养导向下基于ADDIE模型的高中数学单元作业设计探究_杨啸宇.pdf", "education"),
    SourceDoc("基于“教—学—评”一致性理念的“三新”背景下高中数学作业设计研究_华正兰.pdf", "education"),
    SourceDoc("基于核心素养的初中语文写作情境教学研究_滕倩倩.pdf", "education"),
    SourceDoc("小学中段整本书阅读教学策略研究——以X小学为例_程凡芸.pdf", "education"),
    SourceDoc("中国传统文化绘本在小学语文阅读教学中的应用研究_龚婧.pdf", "education"),
    SourceDoc("ESG表现对企业避税的影响研究_常林.pdf", "finance_management"),
    SourceDoc("XH公司绿色债券发行动因与效应研究_吕嘉玮.pdf", "finance_management"),
    SourceDoc("海尔智家数字化转型的财务效果评价研究_李璇.pdf", "finance_management"),
    SourceDoc("工商银行JN分行对公信贷业务发展战略研究_颜寒.pdf", "finance_management"),
    SourceDoc("L公司原材料供应商管理策略优化_冯翊萱.pdf", "finance_management"),
    SourceDoc("数字化转型对工程机械类企业绩效的影响研究_张敬安.pdf", "finance_management"),
    SourceDoc("基于深度学习的CT图像肺部结节检测与分割研究_杜禹博.pdf", "medical_material"),
    SourceDoc("基于频域知识驱动神经网络的多模态医学图像处理方法研究_李哲.pdf", "medical_material"),
    SourceDoc("3D打印海藻酸钠复合水凝胶改性软骨组织的制备与性能研究_孙基桔.pdf", "medical_material"),
    SourceDoc("基于聚乙烯醇水凝胶的功能化应用研究_耿敬天.pdf", "medical_material"),
    SourceDoc("聚乳酸基静电纺丝载药纤维膜的构筑及其药物释放与抗菌性能研究_陈实.pdf", "medical_material"),
    SourceDoc("改性氧化锌催化剂制备及催化降解有机污染物的研究_何书桓.pdf", "medical_material"),
]


DIRECT_SOURCES = [
    "多特征融合的短文本语义相似度计算方法研究_张凝.pdf",
    "基于Transformer的多尺度特征融合的小目标检测研究_王德玉.pdf",
    "C县农产品物流园区选址及功能区布局规划研究_张明慧.pdf",
    "个人数字遗产继承的法律保护研究_李苗苗.pdf",
    "小学中段整本书阅读教学策略研究——以X小学为例_程凡芸.pdf",
    "ESG表现对企业避税的影响研究_常林.pdf",
    "基于深度学习的CT图像肺部结节检测与分割研究_杜禹博.pdf",
    "基于聚乙烯醇水凝胶的功能化应用研究_耿敬天.pdf",
]

REWRITE_SOURCES = [
    "深度学习框架下系统查重的研究与应用_雷子旺.pdf",
    "基于TF-IDF和机器学习的文本向量化与分类研究_石慧.pdf",
    "基于卷积神经网络的人脸表情识别与应用_王以诺.pdf",
    "京津冀农产品物流需求预测与物流网络构建研究_薛晓博.pdf",
    "数字遗产继承问题研究_李欣.pdf",
    "小学中段整本书阅读教学策略研究——以X小学为例_程凡芸.pdf",
    "XH公司绿色债券发行动因与效应研究_吕嘉玮.pdf",
    "改性氧化锌催化剂制备及催化降解有机污染物的研究_何书桓.pdf",
]

SAME_TOPIC_PAIRS = [
    ("深度学习框架下系统查重的研究与应用_雷子旺.pdf", "多特征融合的短文本语义相似度计算方法研究_张凝.pdf"),
    ("基于TF-IDF和机器学习的文本向量化与分类研究_石慧.pdf", "深度学习框架下系统查重的研究与应用_雷子旺.pdf"),
    ("基于Transformer的多尺度特征融合的小目标检测研究_王德玉.pdf", "基于卷积神经网络的人脸表情识别与应用_王以诺.pdf"),
    ("FCH公司生鲜农产品配送路径优化研究_徐小雨.pdf", "双碳目标下生鲜农产品冷链配送路径优化研究_牛忍.pdf"),
    ("平台处罚的行政法规制_吴树贞.pdf", "影评类短视频侵权纠纷中合理使用认定研究_连宇馨.pdf"),
    ("小学中段整本书阅读教学策略研究——以X小学为例_程凡芸.pdf", "中国传统文化绘本在小学语文阅读教学中的应用研究_龚婧.pdf"),
    ("XH公司绿色债券发行动因与效应研究_吕嘉玮.pdf", "ESG表现对企业避税的影响研究_常林.pdf"),
    ("基于聚乙烯醇水凝胶的功能化应用研究_耿敬天.pdf", "改性氧化锌催化剂制备及催化降解有机污染物的研究_何书桓.pdf"),
]

UNRELATED_PAIRS = [
    ("多特征融合的短文本语义相似度计算方法研究_张凝.pdf", "影评类短视频侵权纠纷中合理使用认定研究_连宇馨.pdf"),
    ("基于Transformer的多尺度特征融合的小目标检测研究_王德玉.pdf", "小学中段整本书阅读教学策略研究——以X小学为例_程凡芸.pdf"),
    ("农业全产业链发展水平评价研究_徐靖.pdf", "基于深度学习的CT图像肺部结节检测与分割研究_杜禹博.pdf"),
    ("ESG表现对企业避税的影响研究_常林.pdf", "改性氧化锌催化剂制备及催化降解有机污染物的研究_何书桓.pdf"),
    ("“新事实-新规范”路径下虚拟财产的解释_陈亦洲.pdf", "基于Transformer的多尺度特征融合的小目标检测研究_王德玉.pdf"),
    ("中国传统文化绘本在小学语文阅读教学中的应用研究_龚婧.pdf", "政府规制下生鲜农产品供应链的区块链溯源决策研究_蒋凤娇.pdf"),
]


REPLACEMENTS = [
    ("本文", "本研究"),
    ("文章", "研究"),
    ("研究", "探讨"),
    ("分析", "剖析"),
    ("方法", "方式"),
    ("模型", "框架"),
    ("系统", "平台"),
    ("通过", "采用"),
    ("基于", "依托"),
    ("构建", "建立"),
    ("设计", "规划"),
    ("优化", "改进"),
    ("影响", "作用"),
    ("问题", "现象"),
    ("结果表明", "结果显示"),
    ("具有", "具备"),
    ("提高", "提升"),
    ("降低", "减少"),
    ("应用", "使用"),
    ("实现", "完成"),
    ("提出", "给出"),
    ("首先", "第一"),
    ("其次", "第二"),
    ("最后", "最终"),
]


def cjk_count(text: str) -> int:
    return len(re.findall(r"[\u4e00-\u9fff]", text or ""))


def normalize_text(text: str) -> str:
    text = text.replace("\u3000", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def is_noise_line(line: str) -> bool:
    compact = re.sub(r"\s+", "", line)
    if not compact:
        return True
    if len(compact) <= 6:
        return True
    if re.fullmatch(r"[-_—=·.。0-9A-Za-z]+", compact):
        return True
    if re.match(r"^(分类号|密级|学校代码|学号|作者姓名|导师姓名|单位代码|UDC|U D C)", compact, re.I):
        return True
    if re.match(r"^(摘要|Abstract|目录|参考文献|致谢|攻读|附录)$", compact, re.I):
        return True
    if cjk_count(compact) < 8 and len(compact) < 40:
        return True
    return False


def is_good_paragraph(text: str) -> bool:
    compact = re.sub(r"\s+", "", text)
    if len(compact) < 60:
        return False
    if cjk_count(compact) < 45:
        return False
    digit_ratio = sum(ch.isdigit() for ch in compact) / max(1, len(compact))
    if digit_ratio > 0.25:
        return False
    if compact.count(".") > 12 and cjk_count(compact) / max(1, len(compact)) < 0.45:
        return False
    if any(key in compact[:30] for key in ("目录", "参考文献", "致谢")):
        return False
    return True


def split_sentences(text: str) -> List[str]:
    pieces = re.split(r"(?<=[。！？；;])", text)
    return [piece.strip() for piece in pieces if len(piece.strip()) >= 12]


def extract_pdf_text(path: Path) -> str:
    try:
        import fitz
    except ImportError as exc:
        raise RuntimeError("PyMuPDF is required. Install it with: pip install PyMuPDF") from exc

    paragraphs: List[str] = []
    buffer: List[str] = []

    def flush_buffer(force: bool = False) -> None:
        nonlocal buffer
        if not buffer:
            return
        paragraph = normalize_text("".join(buffer))
        buffer = []
        if force or is_good_paragraph(paragraph):
            if is_good_paragraph(paragraph):
                paragraphs.append(paragraph)

    with fitz.open(path) as doc:
        for page in doc:
            text = page.get_text("text", sort=True)
            lines = [line.strip() for line in text.splitlines()]
            for line in lines:
                if is_noise_line(line):
                    flush_buffer()
                    continue
                compact = re.sub(r"\s+", "", line)
                if not compact:
                    flush_buffer()
                    continue
                if re.match(r"^\d+(?:\.\d+){0,4}\s*[\u4e00-\u9fff]{0,16}$", compact):
                    flush_buffer()
                    continue
                buffer.append(compact)
                joined = "".join(buffer)
                if compact.endswith(("。", "！", "？", "；", ";")) or len(joined) >= 420:
                    flush_buffer()
            flush_buffer()
    flush_buffer(force=True)
    return "\n\n".join(paragraphs)


def extract_docx_text(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    paragraphs = [normalize_text(p.text) for p in doc.paragraphs]
    paragraphs = [p for p in paragraphs if is_good_paragraph(p)]
    return "\n\n".join(paragraphs)


def extract_text(path: Path) -> str:
    try:
        from document_readers.factory import read_document_by_type

        text = read_document_by_type(str(path))
        if text and cjk_count(text) >= 800:
            return text
    except Exception:
        pass

    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf_text(path)
    if suffix == ".docx":
        return extract_docx_text(path)
    if suffix == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore")
    raise ValueError(f"Unsupported file type: {path}")


def required_source_filenames() -> set[str]:
    names = set(DIRECT_SOURCES) | set(REWRITE_SOURCES)
    for target_file, ref_file in SAME_TOPIC_PAIRS + UNRELATED_PAIRS:
        names.add(target_file)
        names.add(ref_file)
    return names


def text_to_paragraphs(text: str) -> List[str]:
    normalized = normalize_text(text)
    normalized = re.sub(r"\s+", "", normalized)
    sentences = split_sentences(normalized)

    filtered: List[str] = []
    for sentence in sentences:
        sentence = normalize_text(sentence)
        compact = re.sub(r"\s+", "", sentence)
        if len(compact) < 18 or len(compact) > 260:
            continue
        if cjk_count(compact) < 12:
            continue
        if is_noise_line(compact):
            continue
        if re.match(r"^(分类号|学号|密级|作者|导师|学院|专业|答辩|摘要|关键词|目录)", compact):
            continue
        if re.match(r"^(第[一二三四五六七八九十0-9]+章|[0-9]+(?:\.[0-9]+){0,4})", compact) and len(compact) < 45:
            continue
        filtered.append(compact)

    paragraphs: List[str] = []
    buffer: List[str] = []
    total = 0
    for sentence in filtered:
        if total + len(sentence) > 360 and total >= 120:
            paragraph = "".join(buffer)
            if is_good_paragraph(paragraph):
                paragraphs.append(paragraph)
            buffer = [sentence]
            total = len(sentence)
        else:
            buffer.append(sentence)
            total += len(sentence)

    if buffer:
        paragraph = "".join(buffer)
        if is_good_paragraph(paragraph):
            paragraphs.append(paragraph)

    return paragraphs


def build_chunk(paragraphs: List[str], offset: int, min_chars: int = 900, max_chars: int = 1450) -> str:
    if not paragraphs:
        raise ValueError("No usable paragraphs")

    start = offset % len(paragraphs)
    ordered = paragraphs[start:] + paragraphs[:start]
    selected: List[str] = []
    total = 0
    for para in ordered:
        selected.append(para)
        total += len(para)
        if total >= min_chars:
            break

    chunk = "\n\n".join(selected)
    if len(chunk) > max_chars:
        sentences = split_sentences(chunk)
        trimmed: List[str] = []
        total = 0
        for sentence in sentences:
            if total + len(sentence) > max_chars and total >= min_chars:
                break
            trimmed.append(sentence)
            total += len(sentence)
        if trimmed:
            chunk = "".join(trimmed)
        chunk = chunk[:max_chars].rstrip("，,；;、")
        if chunk and chunk[-1] not in "。！？":
            chunk += "。"
    return normalize_text(chunk)


def lightly_rewrite(text: str) -> str:
    rewritten = text
    for old, new in REPLACEMENTS:
        rewritten = rewritten.replace(old, new)

    sentences = split_sentences(rewritten)
    if len(sentences) >= 6:
        rearranged: List[str] = []
        for idx in range(0, len(sentences), 4):
            group = sentences[idx : idx + 4]
            if len(group) == 4:
                group = [group[1], group[0], group[2], group[3]]
            rearranged.extend(group)
        sentences = rearranged

    polished: List[str] = []
    prefixes = ["从整体来看，", "进一步而言，", "在实际场景中，", "由此可见，"]
    for idx, sentence in enumerate(sentences):
        sentence = sentence.strip()
        if idx % 5 == 0 and len(sentence) > 25:
            sentence = prefixes[(idx // 5) % len(prefixes)] + sentence
        polished.append(sentence)

    result = "".join(polished)
    result = re.sub(r"。{2,}", "。", result)
    return normalize_text(result)


def domain_of(filename: str, docs: Iterable[SourceDoc]) -> str:
    for doc in docs:
        if doc.filename == filename:
            return doc.domain
    return "unknown"


def load_sources() -> dict[str, ExtractedDoc]:
    loaded: dict[str, ExtractedDoc] = {}
    required_names = required_source_filenames()
    for source in SOURCE_DOCS:
        if source.filename not in required_names:
            continue
        path = DATA_DIR / source.filename
        if not path.exists():
            print(f"[skip] missing: {source.filename}")
            continue
        try:
            text = extract_text(path)
            paragraphs = text_to_paragraphs(text)
            if sum(len(p) for p in paragraphs) < 1500:
                print(f"[skip] too few paragraphs: {source.filename}")
                continue
            char_count = sum(len(p) for p in paragraphs)
            loaded[source.filename] = ExtractedDoc(source, path, paragraphs, char_count)
            print(f"[ok] {source.domain:18s} {source.filename} paragraphs={len(paragraphs)} chars={char_count}")
        except Exception as exc:
            print(f"[skip] failed: {source.filename}: {exc}")
    return loaded


def require_doc(loaded: dict[str, ExtractedDoc], filename: str) -> ExtractedDoc:
    if filename not in loaded:
        raise KeyError(f"Required source was not loaded: {filename}")
    return loaded[filename]


def add_direct_samples(loaded: dict[str, ExtractedDoc], samples: List[SamplePair]) -> None:
    for idx, filename in enumerate(DIRECT_SOURCES, start=1):
        doc = require_doc(loaded, filename)
        chunk = build_chunk(doc.paragraphs, offset=idx * 2)
        samples.append(
            SamplePair(
                case_id=f"C{idx:03d}",
                sample_type="direct_copy",
                label=1,
                target_text=chunk,
                reference_text=chunk,
                target_source=filename,
                reference_source=filename,
                domain=doc.source.domain,
                notes="目标文本与参考文本使用同一来源片段，模拟直接复制。",
            )
        )


def add_rewrite_samples(loaded: dict[str, ExtractedDoc], samples: List[SamplePair], start_index: int) -> None:
    for local_idx, filename in enumerate(REWRITE_SOURCES, start=1):
        doc = require_doc(loaded, filename)
        original = build_chunk(doc.paragraphs, offset=local_idx * 3 + 1)
        rewritten = lightly_rewrite(original)
        samples.append(
            SamplePair(
                case_id=f"C{start_index + local_idx - 1:03d}",
                sample_type="light_rewrite",
                label=1,
                target_text=rewritten,
                reference_text=original,
                target_source=f"{filename} (rewritten)",
                reference_source=filename,
                domain=doc.source.domain,
                notes="目标文本由参考片段进行轻度换词、局部调序和连接词改写，模拟轻度改写。",
            )
        )


def add_same_topic_samples(loaded: dict[str, ExtractedDoc], samples: List[SamplePair], start_index: int) -> None:
    for local_idx, (target_file, ref_file) in enumerate(SAME_TOPIC_PAIRS, start=1):
        target = require_doc(loaded, target_file)
        ref = require_doc(loaded, ref_file)
        target_chunk = build_chunk(target.paragraphs, offset=local_idx * 2 + 4)
        ref_chunk = build_chunk(ref.paragraphs, offset=local_idx * 2 + 7)
        samples.append(
            SamplePair(
                case_id=f"C{start_index + local_idx - 1:03d}",
                sample_type="same_topic_non_plagiarism",
                label=0,
                target_text=target_chunk,
                reference_text=ref_chunk,
                target_source=target_file,
                reference_source=ref_file,
                domain=target.source.domain,
                notes="同一或相近领域的不同文档片段，作为同主题但不应判相似的负样本。",
            )
        )


def add_unrelated_samples(loaded: dict[str, ExtractedDoc], samples: List[SamplePair], start_index: int) -> None:
    for local_idx, (target_file, ref_file) in enumerate(UNRELATED_PAIRS, start=1):
        target = require_doc(loaded, target_file)
        ref = require_doc(loaded, ref_file)
        target_chunk = build_chunk(target.paragraphs, offset=local_idx * 5 + 2)
        ref_chunk = build_chunk(ref.paragraphs, offset=local_idx * 5 + 9)
        samples.append(
            SamplePair(
                case_id=f"C{start_index + local_idx - 1:03d}",
                sample_type="unrelated",
                label=0,
                target_text=target_chunk,
                reference_text=ref_chunk,
                target_source=target_file,
                reference_source=ref_file,
                domain=f"{target.source.domain} vs {ref.source.domain}",
                notes="跨领域文档片段，作为完全无关负样本。",
            )
        )


def summarize_text(text: str) -> dict[str, int]:
    sentences = split_sentences(text)
    return {
        "chars": len(text),
        "cjk_chars": cjk_count(text),
        "sentences": len(sentences),
    }


def write_sample_files(samples: List[SamplePair]) -> None:
    if OUTPUT_DIR.exists():
        shutil.rmtree(OUTPUT_DIR)
    TARGET_DIR.mkdir(parents=True, exist_ok=True)
    REFERENCE_DIR.mkdir(parents=True, exist_ok=True)

    rows = []
    json_rows = []
    for sample in samples:
        target_path = TARGET_DIR / f"{sample.case_id}_target.txt"
        ref_path = REFERENCE_DIR / f"{sample.case_id}_ref.txt"
        target_path.write_text(sample.target_text + "\n", encoding="utf-8")
        ref_path.write_text(sample.reference_text + "\n", encoding="utf-8")

        target_summary = summarize_text(sample.target_text)
        ref_summary = summarize_text(sample.reference_text)
        row = {
            "case_id": sample.case_id,
            "sample_type": sample.sample_type,
            "label": sample.label,
            "target_file": str(target_path.relative_to(OUTPUT_DIR)).replace("\\", "/"),
            "reference_file": str(ref_path.relative_to(OUTPUT_DIR)).replace("\\", "/"),
            "target_source": sample.target_source,
            "reference_source": sample.reference_source,
            "domain": sample.domain,
            "target_chars": target_summary["chars"],
            "reference_chars": ref_summary["chars"],
            "target_sentences": target_summary["sentences"],
            "reference_sentences": ref_summary["sentences"],
            "notes": sample.notes,
        }
        rows.append(row)
        json_rows.append(row)

    with (OUTPUT_DIR / "labels.csv").open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=list(rows[0].keys()))
        writer.writeheader()
        writer.writerows(rows)

    with (OUTPUT_DIR / "labels.jsonl").open("w", encoding="utf-8") as f:
        for row in json_rows:
            f.write(json.dumps(row, ensure_ascii=False) + "\n")


def write_file_source_table(samples: List[SamplePair]) -> None:
    rows = []
    for sample in samples:
        target_summary = summarize_text(sample.target_text)
        reference_summary = summarize_text(sample.reference_text)
        target_file = f"samples/targets/{sample.case_id}_target.txt"
        reference_file = f"samples/references/{sample.case_id}_ref.txt"

        rows.append(
            {
                "file_path": target_file,
                "case_id": sample.case_id,
                "role": "target",
                "paired_file": reference_file,
                "sample_type": sample.sample_type,
                "label": sample.label,
                "source_document": sample.target_source,
                "domain": sample.domain,
                "chars": target_summary["chars"],
                "sentences": target_summary["sentences"],
                "notes": sample.notes,
            }
        )
        rows.append(
            {
                "file_path": reference_file,
                "case_id": sample.case_id,
                "role": "reference",
                "paired_file": target_file,
                "sample_type": sample.sample_type,
                "label": sample.label,
                "source_document": sample.reference_source,
                "domain": sample.domain,
                "chars": reference_summary["chars"],
                "sentences": reference_summary["sentences"],
                "notes": sample.notes,
            }
        )

    with (OUTPUT_DIR / "sample_file_sources.csv").open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=list(rows[0].keys()))
        writer.writeheader()
        writer.writerows(rows)

    lines = [
        "# \u6837\u672c\u6587\u4ef6\u6765\u6e90\u8868",
        "",
        "\u672c\u8868\u9010\u4e00\u8bf4\u660e 60 \u4e2a txt \u6837\u672c\u6587\u4ef6\u7684\u6765\u6e90\u6587\u6863\u3001\u6240\u5c5e\u6837\u672c\u5bf9\u3001\u89d2\u8272\u548c\u6807\u7b7e\u3002",
        "",
        "| file_path | case_id | role | sample_type | label | source_document | domain | chars |",
        "|---|---|---|---|---:|---|---|---:|",
    ]
    for row in rows:
        lines.append(
            "| {file_path} | {case_id} | {role} | {sample_type} | {label} | {source_document} | {domain} | {chars} |".format(
                **row
            )
        )
    (OUTPUT_DIR / "sample_file_sources.md").write_text("\n".join(lines) + "\n", encoding="utf-8-sig")


def write_source_manifest(loaded: dict[str, ExtractedDoc]) -> None:
    rows = []
    for filename, doc in sorted(loaded.items(), key=lambda item: (item[1].source.domain, item[0])):
        rows.append(
            {
                "filename": filename,
                "domain": doc.source.domain,
                "paragraph_count": len(doc.paragraphs),
                "char_count": doc.char_count,
                "used_as_configured_source": "yes",
            }
        )
    with (OUTPUT_DIR / "source_manifest.csv").open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=list(rows[0].keys()))
        writer.writeheader()
        writer.writerows(rows)


def write_threshold_grid() -> None:
    rows = []
    for value in [0.04, 0.06, 0.08, 0.10, 0.12, 0.15]:
        rows.append(
            {
                "experiment_stage": "window_threshold",
                "window_threshold": value,
                "fine_threshold": 0.30,
                "description": "固定细粒度阈值，观察窗口命中阈值变化。",
            }
        )
    for value in [0.20, 0.25, 0.30, 0.35, 0.40]:
        rows.append(
            {
                "experiment_stage": "fine_threshold",
                "window_threshold": 0.08,
                "fine_threshold": value,
                "description": "固定窗口阈值，观察细粒度确认阈值变化。",
            }
        )
    with (OUTPUT_DIR / "threshold_grid.csv").open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=list(rows[0].keys()))
        writer.writeheader()
        writer.writerows(rows)


def write_readme(samples: List[SamplePair]) -> None:
    type_counts: dict[str, int] = {}
    label_counts: dict[int, int] = {}
    for sample in samples:
        type_counts[sample.sample_type] = type_counts.get(sample.sample_type, 0) + 1
        label_counts[sample.label] = label_counts.get(sample.label, 0) + 1

    readme = f"""# 阈值敏感性实验样本集

本目录保存传统模式阈值敏感性实验使用的样本文档对。样本由 `data/` 目录中的论文文档抽取并构造，分为直接复制、轻度改写、同主题不抄袭和完全无关四类。

## 文件说明

- `samples/targets/`：待检测文本。
- `samples/references/`：参考文本。
- `labels.csv`：样本标签、来源文档、样本类型和长度信息。
- `labels.jsonl`：与 `labels.csv` 等价的 JSON Lines 格式。
- `threshold_grid.csv`：建议测试的阈值组合。
- `source_manifest.csv`：成功解析并用于样本构造的源文档清单。

## 标签含义

- `label = 1`：应判为相似。
- `label = 0`：不应判为相似。

## 样本构成

样本总数：{len(samples)}

| sample_type | count |
|---|---:|
{chr(10).join(f'| {key} | {value} |' for key, value in sorted(type_counts.items()))}

| label | count |
|---|---:|
{chr(10).join(f'| {key} | {value} |' for key, value in sorted(label_counts.items()))}

## 建议实验流程

1. 第一轮固定细粒度阈值为 `0.30`，测试 `threshold_grid.csv` 中的窗口阈值。
2. 根据 Precision、Recall、F1、FP、FN 选择较均衡的窗口阈值。
3. 第二轮固定选出的窗口阈值，再测试细粒度确认阈值。
4. 使用 `labels.csv` 作为真实标签表。

第一次统计时，可以先把 `hit_count > 0` 定义为系统检出；如果后续需要更严格，可以改成 `hit_count > 0 and score >= fine_threshold`。
"""
    (OUTPUT_DIR / "README.md").write_text(readme, encoding="utf-8-sig")


def validate_samples(samples: List[SamplePair]) -> None:
    if len(samples) != 30:
        raise AssertionError(f"Expected 30 samples, got {len(samples)}")
    ids = [sample.case_id for sample in samples]
    if len(set(ids)) != len(ids):
        raise AssertionError("Duplicate case_id detected")
    for sample in samples:
        for role, text in [("target", sample.target_text), ("reference", sample.reference_text)]:
            if len(text) < 650:
                raise AssertionError(f"{sample.case_id} {role} is too short: {len(text)}")
            if cjk_count(text) < 450:
                raise AssertionError(f"{sample.case_id} {role} has too few CJK chars")
    positives = sum(1 for sample in samples if sample.label == 1)
    negatives = sum(1 for sample in samples if sample.label == 0)
    if positives != 16 or negatives != 14:
        raise AssertionError(f"Expected 16 positive and 14 negative samples, got {positives}/{negatives}")


def main() -> None:
    loaded = load_sources()
    samples: List[SamplePair] = []
    add_direct_samples(loaded, samples)
    add_rewrite_samples(loaded, samples, start_index=len(samples) + 1)
    add_same_topic_samples(loaded, samples, start_index=len(samples) + 1)
    add_unrelated_samples(loaded, samples, start_index=len(samples) + 1)
    validate_samples(samples)
    write_sample_files(samples)
    write_file_source_table(samples)
    write_source_manifest(loaded)
    write_threshold_grid()
    write_readme(samples)

    print()
    print(f"Wrote {len(samples)} sample pairs to: {OUTPUT_DIR}")
    print(f"Positive pairs: {sum(1 for sample in samples if sample.label == 1)}")
    print(f"Negative pairs: {sum(1 for sample in samples if sample.label == 0)}")


if __name__ == "__main__":
    main()
